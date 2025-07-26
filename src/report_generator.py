from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_rfe_risk_report(analyses, output_filename="RFE_Risk_Report.docx"):
    """
    Generates a professional DOCX report from the collected AI analyses.
    'analyses' is a dictionary where keys are segment headers and values are the AI's markdown analysis.
    """
    print("Generating final DOCX report...")
    doc = Document()

    # Title Page
    doc.add_heading('VisaCompanion RFE Risk Analyzer', level=0)
    doc.add_heading('EB-1A Petition Analysis Report', level=1)
    p = doc.add_paragraph('Confidential Internal Memo')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        "This report provides a detailed analysis of a draft EB-1A petition, identifying potential "
        "weaknesses that may trigger a Request for Evidence (RFE) from USCIS. Each section of the "
        "petition has been evaluated against the corresponding EB-1A criteria, with specific, "
        "actionable recommendations provided to strengthen the case."
    )

    # Analysis Sections
    for header, markdown_analysis in analyses.items():
        doc.add_heading(f"Analysis of Section: {header}", level=2)

        lines = markdown_analysis.strip().replace("```markdown", "").replace("```", "").strip().split('\n')

        in_table = False
        table = None
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith('|'):
                if not in_table:
                    num_cols = len(line.split('|')) - 2
                    table = doc.add_table(rows=1, cols=num_cols)
                    table.style = 'Table Grid'
                    in_table = True
                    header_cells = table.rows[0].cells
                    # Set table headers
                    headers = [h.strip() for h in line.strip('|').split('|')]
                    for i, h in enumerate(headers):
                        header_cells[i].text = h
                else:
                    # Add row to table (skipping header separator)
                    if '---' not in line:
                        row_cells = table.add_row().cells
                        cells = [c.strip() for c in line.strip('|').split('|')]
                        for i, c in enumerate(cells):
                            row_cells[i].text = c.replace('**', '')
            else:
                in_table = False
                if line.startswith('**'):
                    doc.add_paragraph(line.replace('**', ''), style='Intense Quote')
                else:
                    doc.add_paragraph(line)
        doc.add_paragraph()
    
    doc.save(output_filename)
    print(f"✅ Report successfully saved as {output_filename}")